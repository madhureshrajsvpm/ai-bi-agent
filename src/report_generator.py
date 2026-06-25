from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd
import plotly.express as px
import plotly.io as pio
import io
import os
from datetime import datetime


def set_cell_bg(cell, hex_color):
    """Set background color of a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1, color="1F3864"):
    """Add a styled heading."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(color)
    if level == 1:
        run.font.size = Pt(18)
    elif level == 2:
        run.font.size = Pt(14)
    else:
        run.font.size = Pt(12)
    return p


def add_divider(doc, color="2E75B6"):
    """Add a horizontal line."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def generate_chart_image(df, x_col, y_col, chart_title):
    """Generate a Plotly chart and return as bytes."""
    try:
        num_cols = df.select_dtypes(include="number").columns.tolist()
        cat_cols = df.select_dtypes(include="object").columns.tolist()

        if not x_col or not y_col:
            return None
        if x_col not in df.columns or y_col not in df.columns:
            return None

        agg = (
            df.groupby(x_col)[y_col]
            .sum().reset_index()
            .sort_values(y_col, ascending=False)
            .head(10)
        )

        fig = px.bar(
            agg, x=x_col, y=y_col,
            title=chart_title,
            color=y_col,
            color_continuous_scale="Blues"
        )
        fig.update_layout(
            width=700, height=400,
            paper_bgcolor="white",
            plot_bgcolor="white",
            font=dict(family="Arial", size=12)
        )

        img_bytes = pio.to_image(fig, format="png", width=700, height=400)
        return img_bytes
    except Exception:
        return None


def generate_report(
    df,
    dataset_name,
    insights_text,
    profile,
    author_name="Madhuresh Raj Selvaraj"
):
    """Generate a Word report and return as bytes."""
    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)

    # ── Cover header ──────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("AI Business Intelligence Report")
    title_run.bold = True
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor.from_string("1F3864")

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_p.add_run(dataset_name)
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = RGBColor.from_string("2E75B6")

    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta_p.add_run(
        f"Prepared by: {author_name}   |   Generated: {datetime.now().strftime('%B %d, %Y')}"
    )
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor.from_string("666666")

    add_divider(doc)
    doc.add_paragraph()

    # ── Dataset Overview ──────────────────────────────────────
    add_heading(doc, "1. Dataset Overview", level=1)

    table = doc.add_table(rows=2, cols=4)
    table.style = "Table Grid"

    headers = ["Total Rows", "Total Columns", "Null Values", "Duplicate Rows"]
    values  = [
        f"{profile.get('total_rows', 0):,}",
        str(profile.get('total_columns', 0)),
        str(sum(v['count'] for v in profile.get('null_summary', {}).values())),
        str(profile.get('duplicate_rows', 0))
    ]

    header_row = table.rows[0]
    value_row  = table.rows[1]

    for i, (h, v) in enumerate(zip(headers, values)):
        hc = header_row.cells[i]
        vc = value_row.cells[i]

        set_cell_bg(hc, "1F3864")
        hc.paragraphs[0].clear()
        hr = hc.paragraphs[0].add_run(h)
        hr.bold = True
        hr.font.color.rgb = RGBColor(255, 255, 255)
        hr.font.size = Pt(10)
        hc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        vc.paragraphs[0].clear()
        vr = vc.paragraphs[0].add_run(v)
        vr.bold = True
        vr.font.size = Pt(14)
        vc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Column list
    add_heading(doc, "Columns", level=3)
    cols_text = ", ".join(profile.get("column_names", []))
    p = doc.add_paragraph(cols_text)
    p.runs[0].font.size = Pt(10)

    # Null summary
    null_summary = profile.get("null_summary", {})
    if null_summary:
        doc.add_paragraph()
        add_heading(doc, "Data Quality — Null Values", level=2)
        null_table = doc.add_table(rows=1, cols=3)
        null_table.style = "Table Grid"

        for i, header in enumerate(["Column", "Null Count", "% of Rows"]):
            cell = null_table.rows[0].cells[i]
            set_cell_bg(cell, "2E75B6")
            cell.paragraphs[0].clear()
            r = cell.paragraphs[0].add_run(header)
            r.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.font.size = Pt(10)

        for col, info in null_summary.items():
            row = null_table.add_row()
            row.cells[0].paragraphs[0].add_run(col).font.size = Pt(10)
            row.cells[1].paragraphs[0].add_run(str(info['count'])).font.size = Pt(10)
            row.cells[2].paragraphs[0].add_run(f"{info['pct']}%").font.size = Pt(10)

    doc.add_paragraph()
    add_divider(doc)

    # ── Auto Insights ──────────────────────────────────────────
    doc.add_paragraph()
    add_heading(doc, "2. Key Business Insights", level=1)

    if insights_text:
        for line in insights_text.strip().split("\n"):
            line = line.strip()
            if not line:
                continue
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(line)
            run.font.size = Pt(11)
            if line.startswith("⚠️"):
                run.font.color.rgb = RGBColor.from_string("C00000")
            elif line.startswith("🏆"):
                run.font.color.rgb = RGBColor.from_string("375623")
            elif line.startswith("🔗"):
                run.font.color.rgb = RGBColor.from_string("1F3864")
            elif line.startswith("💡"):
                run.font.color.rgb = RGBColor.from_string("7030A0")
            else:
                run.font.color.rgb = RGBColor.from_string("000000")

    doc.add_paragraph()
    add_divider(doc)

    # ── Numeric Stats ──────────────────────────────────────────
    numeric_stats = profile.get("numeric_stats", {})
    if numeric_stats:
        doc.add_paragraph()
        add_heading(doc, "3. Numeric Column Statistics", level=1)

        stat_cols = ["Column", "Min", "Max", "Mean", "Median", "Std Dev"]
        stat_table = doc.add_table(rows=1, cols=len(stat_cols))
        stat_table.style = "Table Grid"

        for i, header in enumerate(stat_cols):
            cell = stat_table.rows[0].cells[i]
            set_cell_bg(cell, "1F3864")
            cell.paragraphs[0].clear()
            r = cell.paragraphs[0].add_run(header)
            r.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.font.size = Pt(9)

        for col, stats in numeric_stats.items():
            row = stat_table.add_row()
            vals = [
                col,
                str(stats.get("min", "")),
                str(stats.get("max", "")),
                str(stats.get("mean", "")),
                str(stats.get("median", "")),
                str(stats.get("std", ""))
            ]
            for i, val in enumerate(vals):
                row.cells[i].paragraphs[0].add_run(val).font.size = Pt(9)

        doc.add_paragraph()
        add_divider(doc)

    # ── Correlations ──────────────────────────────────────────
    correlations = profile.get("strong_correlations", [])
    if correlations:
        doc.add_paragraph()
        add_heading(doc, "4. Strong Correlations", level=1)

        corr_table = doc.add_table(rows=1, cols=3)
        corr_table.style = "Table Grid"

        for i, header in enumerate(["Column A", "Column B", "Correlation"]):
            cell = corr_table.rows[0].cells[i]
            set_cell_bg(cell, "2E75B6")
            cell.paragraphs[0].clear()
            r = cell.paragraphs[0].add_run(header)
            r.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.font.size = Pt(10)

        for pair in correlations:
            row = corr_table.add_row()
            row.cells[0].paragraphs[0].add_run(pair["col_a"]).font.size = Pt(10)
            row.cells[1].paragraphs[0].add_run(pair["col_b"]).font.size = Pt(10)
            row.cells[2].paragraphs[0].add_run(str(pair["correlation"])).font.size = Pt(10)

        doc.add_paragraph()
        add_divider(doc)

    # ── Chart ──────────────────────────────────────────────────
    num_cols = df.select_dtypes(include="number").columns.tolist()
    cat_cols = df.select_dtypes(include="object").columns.tolist()

    if cat_cols and num_cols:
        doc.add_paragraph()
        add_heading(doc, "5. Top Category Chart", level=1)

        x_col = cat_cols[0]
        y_col = num_cols[0]
        chart_bytes = generate_chart_image(
            df, x_col, y_col,
            f"Top 10: {y_col} by {x_col}"
        )
        if chart_bytes:
            img_stream = io.BytesIO(chart_bytes)
            doc.add_picture(img_stream, width=Inches(6))
            last_p = doc.paragraphs[-1]
            last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()
        add_divider(doc)

    # ── Footer note ───────────────────────────────────────────
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer_p.add_run(
        f"Generated by AI BI Agent · {author_name} · {datetime.now().strftime('%B %d, %Y')}"
    )
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor.from_string("999999")
    fr.italic = True

    # ── Save to bytes ──────────────────────────────────────────
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()